# ==============================================================================
#      FULL ENHANCED STREAMLIT APPLICATION (with Feature-Rich Sidebar)
# ==============================================================================

# ==============================================================================
# 1. IMPORTS
# ==============================================================================
import streamlit as st
import tensorflow as tf
from tensorflow.keras.models import Model, load_model
from tensorflow.keras.applications import efficientnet_v2
from tensorflow.keras.layers import Conv2D, SeparableConv2D, DepthwiseConv2D
from PIL import Image, ImageDraw
import numpy as np
import pandas as pd
import datetime
import os
import io
import cv2

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==============================================================================
# 2. CONFIGURATION
# ==============================================================================
class Config:
    """Holds all static configuration for the app."""
    APP_TITLE = "NeuroScan AI – Revolutionizing Brain Tumor Analysis"
    PAGE_ICON = "🧠"
    CLASS_NAMES = ['Glioma Tumor', 'Meningioma Tumor', 'No Tumor', 'Pituitary Tumor']
    IMG_SIZE = 224

    BASE_DIR = os.path.dirname(os.path.abspath(__file__))
    CLASSIFICATION_MODEL_PATH = os.path.join(BASE_DIR, 'best_brain_tumor_model_efficient_refined.h5')

if not os.path.exists(Config.CLASSIFICATION_MODEL_PATH):
    st.error(
        f"FATAL ERROR: Model file not found at path: {Config.CLASSIFICATION_MODEL_PATH}\n"
        "Please ensure 'best_brain_tumor_model_efficient_refined.h5' is in the same directory as this script."
    )
    st.stop()


# ==============================================================================
# 3. THEME & STYLING
# ==============================================================================
def load_css():
    """Injects custom CSS for a modern, professional dark theme."""
    st.markdown("""
    <style>
        /* (CSS is unchanged) */
        :root {
            --bg-dark: #0a0a0a; --bg-light: #1a1a1a; --border: #333333; --accent-primary: #0078f2;
            --accent-secondary: #00ab66; --text-primary: #e6e6e6; --text-secondary: #a0a0a0;
            --highlight-bad: #d93025; --highlight-good: #1e8e3e; --shadow-color: rgba(0, 0, 0, 0.4);
        }
        html, body, .stApp { background-color: var(--bg-dark); color: var(--text-primary); }
        h1, h2, h3 { color: var(--text-primary); }
        .main .stMarkdown h1 { text-align: center; font-weight: 700; }
        [data-testid="stSidebar"] { background-color: var(--bg-light); border-right: 1px solid var(--border); }
        .stButton>button { background-color: var(--accent-primary); color: #FFFFFF; font-weight: bold; border-radius: 8px; border: none; box-shadow: 0 4px 10px var(--shadow-color); transition: background-color 0.2s ease; }
        .stButton>button:hover { background-color: #005bb5; }
        .stTextInput>div>div>input, .stSelectbox>div>div, .stNumberInput>div>div>input { background-color: var(--bg-dark); border: 1px solid var(--border); border-radius: 8px; }
        .stFileUploader div[data-testid="stFileUploaderDropzone"] { background-color: var(--bg-dark); border: 2px dashed var(--border); border-radius: 8px; }
        .highlight-diagnosis { background-color: var(--highlight-bad); color: white; padding: 6px 12px; border-radius: 8px; font-weight: bold; }
        .highlight-no-tumor { background-color: var(--highlight-good); color: white; padding: 6px 12px; border-radius: 8px; font-weight: bold; }
        .stTabs [data-baseweb="tab"] { background-color: transparent; border-bottom: 2px solid transparent; }
        .stTabs [aria-selected="true"] { color: var(--accent-primary); border-bottom: 2px solid var(--accent-primary); }
        .custom-section-container, .report-container { background-color: var(--bg-light); border: 1px solid var(--border); border-radius: 12px; padding: 2rem; margin-bottom: 2rem; box-shadow: 0 8px 20px var(--shadow-color); }
        .custom-section-header { color: var(--text-primary); font-size: 1.8em; border-bottom: 1px solid var(--border); padding-bottom: 0.5rem; margin-bottom: 1rem; display: flex; align-items: center; gap: 0.5rem; }
        .report-header { text-align: center; color: var(--accent-primary); font-size: 2.2em; border-bottom: 1px solid var(--border); padding-bottom: 15px; margin-bottom: 25px; }
        .report-section { background-color: var(--bg-dark); border-radius: 10px; padding: 20px; margin-bottom: 20px; border: 1px solid var(--border); }
        .report-section-title { color: var(--text-primary); font-size: 1.4em; margin-bottom: 15px; }
        .report-item { display: flex; justify-content: space-between; align-items: center; padding: 10px 0; border-bottom: 1px solid var(--border); }
        .report-item:last-child { border-bottom: none; }
        .report-item-label { color: var(--text-secondary); }
        .report-item-value { color: var(--text-primary); font-weight: 500; }
        .highlight-confidence { font-weight: bold; color: var(--accent-secondary); }
        .report-disclaimer { font-size: 0.9em; color: var(--text-secondary); text-align: justify; margin-top: 20px; padding: 15px; background-color: var(--bg-dark); border-radius: 8px; border: 1px dashed var(--border); }
    </style>
    """, unsafe_allow_html=True)

# ==============================================================================
# 4. CORE AI & LOGIC FUNCTIONS
# ==============================================================================
@st.cache_resource
def load_keras_model(model_path):
    """Load the Keras model from disk, with caching."""
    try:
        return load_model(model_path)
    except Exception as e:
        st.error(f"Error loading model: {e}")
        return None

def preprocess_image(image):
    """Preprocess the uploaded image for the classification model."""
    image = image.resize((Config.IMG_SIZE, Config.IMG_SIZE))
    image_array = np.array(image)
    if image_array.ndim == 2: image_array = np.stack((image_array,) * 3, axis=-1)
    if image_array.shape[2] == 4: image_array = image_array[:, :, :3]
    image_array = np.expand_dims(image_array, axis=0)
    return efficientnet_v2.preprocess_input(image_array)

def generate_grad_cam(model, img_array, original_image_pil, predicted_class_index):
    """Generates and overlays a Grad-CAM heatmap on an image."""
    last_conv_layer_name = None
    for layer in reversed(model.layers):
        if isinstance(layer, (Conv2D, SeparableConv2D, DepthwiseConv2D)):
            last_conv_layer_name = layer.name
            break
    if last_conv_layer_name is None:
        return np.array(original_image_pil)

    grad_model = Model(model.inputs, [model.get_layer(last_conv_layer_name).output, model.output])
    with tf.GradientTape() as tape:
        conv_outputs, preds = grad_model(img_array)
        if isinstance(preds, list): preds = preds[0]
        loss = preds[:, predicted_class_index]

    grads = tape.gradient(loss, conv_outputs)
    if grads is None: return np.array(original_image_pil)

    pooled_grads = tf.reduce_mean(grads, axis=(0, 1, 2))
    heatmap = tf.reduce_mean(tf.multiply(pooled_grads, conv_outputs[0]), axis=-1)
    heatmap = np.maximum(heatmap, 0)
    if np.max(heatmap) > 0: heatmap /= np.max(heatmap)
    heatmap_resized = cv2.resize(heatmap, (Config.IMG_SIZE, Config.IMG_SIZE))
    heatmap_uint8 = np.uint8(255 * heatmap_resized)
    heatmap_color = cv2.applyColorMap(heatmap_uint8, cv2.COLORMAP_JET)
    original_cv_image = cv2.cvtColor(np.array(original_image_pil.resize((Config.IMG_SIZE, Config.IMG_SIZE))), cv2.COLOR_RGB2BGR)
    superimposed_img = cv2.addWeighted(original_cv_image, 0.6, heatmap_color, 0.4, 0)
    superimposed_img = cv2.cvtColor(superimposed_img, cv2.COLOR_BGR2RGB)
    return superimposed_img

def mock_localization(image, class_name):
    # (This function is unchanged)
    width, height = image.size
    overlay = Image.new('RGBA', image.size, (0, 0, 0, 0))
    draw = ImageDraw.Draw(overlay)
    tumor_detected = False
    if class_name != 'No Tumor' and np.random.rand() > 0.1:
        tumor_detected = True
        center_x, center_y = width * np.random.uniform(0.3, 0.7), height * np.random.uniform(0.3, 0.7)
        major_axis, minor_axis = width * np.random.uniform(0.15, 0.35), height * np.random.uniform(0.1, 0.25)
        bbox = [center_x - major_axis/2, center_y - minor_axis/2, center_x + major_axis/2, center_y + minor_axis/2]
        draw.ellipse(bbox, fill=(218, 54, 51, 100), outline=(218, 54, 51, 200), width=3)
    result_image = Image.alpha_composite(image.convert("RGBA"), overlay)
    return result_image.convert("RGB"), tumor_detected

def get_tumor_details(class_name):
    # (This function is unchanged)
    details = {'stage': "Requires clinical assessment.", 'precautions': "Consult a medical professional for a definitive diagnosis.", 'suggestions': "Further imaging and biopsy may be required."}
    if "Glioma" in class_name: details.update({'stage': "Potentially aggressive", 'precautions': "Urgent neurological consultation.", 'suggestions': "Oncological review, surgery, radiation, and/or chemotherapy."})
    elif "Meningioma" in class_name: details.update({'stage': "Often benign", 'precautions': "Regular follow-up MRI scans.", 'suggestions': "Observation or surgical resection."})
    elif "Pituitary" in class_name: details.update({'stage': "Mostly benign", 'precautions': "Endocrinological evaluation.", 'suggestions': "Medical management or surgery."})
    elif "No Tumor" in class_name: details.update({'stage': "N/A", 'precautions': "Continue routine health check-ups.", 'suggestions': "No immediate medical intervention suggested."})
    return details

# ==============================================================================
# 5. DOCX & REPORT GENERATION
# ==============================================================================
def create_doc_report(patient_info, classification_result, tumor_details, tumor_detected):
    # (This function is unchanged)
    document = Document()
    def add_info_row(label, value, bold_value=False, color_rgb=None):
        p = document.add_paragraph(); p.add_run(f'{label}: ').bold = True; value_run = p.add_run(str(value))
        if bold_value: value_run.bold = True
        if color_rgb: value_run.font.color.rgb = color_rgb
    header = document.add_paragraph(); header_run = header.add_run(Config.APP_TITLE); header_run.bold = True; header_run.font.size = Pt(20); header_run.font.color.rgb = RGBColor(0x00, 0x78, 0xf2); header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_date = document.add_paragraph(f"Report Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"); p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_heading('👤 Patient Details', level=1)
    for key, value in patient_info.items(): add_info_row(key, value)
    document.add_heading('🔬 AI Analysis', level=1)
    is_tumor = "No Tumor" not in classification_result['class_name']; diagnosis_color = RGBColor(0xd9, 0x30, 0x25) if is_tumor else RGBColor(0x1e, 0x8e, 0x3e)
    add_info_row("Classification", classification_result['class_name'], bold_value=True, color_rgb=diagnosis_color)
    add_info_row("Confidence Score", f"{classification_result['confidence']:.2%}", bold_value=True, color_rgb=RGBColor(0x00, 0xab, 0x66))
    add_info_row("Tumor Localization", "Detected" if tumor_detected else "Not Detected")
    document.add_heading('🩺 Provisional Staging & Recommendations', level=1)
    add_info_row("AI Stage Estimate", tumor_details['stage']); add_info_row("Precautions", tumor_details['precautions']); add_info_row("Suggestions", tumor_details['suggestions'])
    document.add_section(); disclaimer = document.add_paragraph(); disclaimer_run = disclaimer.add_run("Disclaimer: This report is generated by an AI model for informational purposes only and is not a substitute for professional medical advice, diagnosis, or treatment."); disclaimer_run.italic = True; disclaimer_run.font.size = Pt(9)
    doc_io = io.BytesIO(); document.save(doc_io); doc_io.seek(0)
    return doc_io.getvalue()

def generate_html_report(patient_info, classification_result, tumor_details, tumor_detected):
    # (This function is unchanged)
    is_tumor = "No Tumor" not in classification_result['class_name']; diagnosis_class = "highlight-diagnosis" if is_tumor else "highlight-no-tumor"
    return f"""<div class="report-container"><div class="report-header">{Config.APP_TITLE}</div><div class="report-section"><div class="report-section-title">👤 Patient Details</div><div class="report-item"><span class="report-item-label">Name:</span> <span class="report-item-value">{patient_info['Name']}</span></div><div class="report-item"><span class="report-item-label">Age:</span> <span class="report-item-value">{patient_info['Age']}</span></div><div class="report-item"><span class="report-item-label">Gender:</span> <span class="report-item-value">{patient_info['Gender']}</span></div><div class="report-item"><span class="report-item-label">Patient ID:</span> <span class="report-item-value">{patient_info['Patient ID']}</span></div></div><div class="report-section"><div class="report-section-title">🔬 AI Analysis</div><div class="report-item"><span class="report-item-label">Classification:</span> <span class="{diagnosis_class}">{classification_result['class_name']}</span></div><div class="report-item"><span class="report-item-label">Confidence Score:</span> <span class="highlight-confidence">{classification_result['confidence']:.2%}</span></div><div class="report-item"><span class="report-item-label">Tumor Localization:</span> <span class="report-item-value">{"Detected" if tumor_detected else "Not Detected"}</span></div></div><div class="report-section"><div class="report-section-title">🩺 Provisional Staging & Recommendations</div><div class="report-item"><span class="report-item-label">AI Stage Estimate:</span> <span class="report-item-value">{tumor_details['stage']}</span></div><div class="report-item"><span class="report-item-label">Precautions:</span> <span class="report-item-value">{tumor_details['precautions']}</span></div><div class="report-item"><span class="report-item-label">Suggestions:</span> <span class="report-item-value">{tumor_details['suggestions']}</span></div></div><div class="report-disclaimer"><strong>Disclaimer:</strong> This report is AI-generated and for informational purposes only. It is not a substitute for professional medical advice.</div></div>"""

# ==============================================================================
# 6. UI VIEW FUNCTIONS
# ==============================================================================
def render_analysis_page():
    # (This function is unchanged)
    with st.expander("👤 Enter Patient Information & Upload MRI", expanded=True):
        with st.form("patient_form"):
            col1, col2 = st.columns(2)
            patient_name = col1.text_input("Patient Name", placeholder="e.g., John Doe")
            patient_id = col2.text_input("Patient ID", placeholder="e.g., P12345")
            patient_age = col1.number_input("Age", min_value=0, max_value=120, value=None, placeholder="e.g., 45")
            patient_gender = col2.selectbox("Gender", [None, "Male", "Female", "Other"])
            uploaded_file = st.file_uploader("Upload Brain MRI Image...", type=["jpg", "jpeg", "png"])
            submit_button = st.form_submit_button(label="🚀 Analyze MRI")

    if submit_button and uploaded_file:
        if not all([patient_name, patient_id, patient_age, patient_gender]):
            st.error("❗ Please fill in all patient details before analyzing.")
            return
        handle_analysis(patient_id, patient_name, patient_age, patient_gender, uploaded_file)

def handle_analysis(patient_id, patient_name, patient_age, patient_gender, uploaded_file):
    st.markdown('<div class="custom-section-container">', unsafe_allow_html=True)
    st.markdown('<h2 class="custom-section-header">🔬 Analysis Results</h2>', unsafe_allow_html=True)
    model = load_keras_model(Config.CLASSIFICATION_MODEL_PATH)
    if model is None: return
    original_image = Image.open(uploaded_file).convert("RGB")
    with st.spinner('Performing AI analysis...'):
        preprocessed_img = preprocess_image(original_image)
        prediction = model.predict(preprocessed_img)
        class_index = np.argmax(prediction, axis=1)[0]
        class_name = Config.CLASS_NAMES[class_index]
        confidence = np.max(prediction)
        classification_result = {"class_name": class_name, "confidence": confidence}
        localized_image, tumor_detected = mock_localization(original_image.copy(), class_name)
        grad_cam_image = generate_grad_cam(model, preprocessed_img, original_image, class_index)
        tumor_details = get_tumor_details(class_name)
    col1, col2, col3 = st.columns(3)
    # --- CORRECTED CODE HERE ---
    col1.image(original_image, caption='Original MRI Scan', use_container_width=True)
    col2.image(localized_image, caption='Mock Tumor Localization', use_container_width=True)
    col3.image(grad_cam_image, caption='Grad-CAM Heatmap', use_container_width=True)
    # --- END OF CORRECTION ---
    st.markdown("<br>", unsafe_allow_html=True)
    patient_info = {"Name": patient_name, "Age": str(patient_age), "Gender": patient_gender, "Patient ID": patient_id}
    st.markdown(generate_html_report(patient_info, classification_result, tumor_details, tumor_detected), unsafe_allow_html=True)
    doc_bytes = create_doc_report(patient_info, classification_result, tumor_details, tumor_detected)
    st.download_button(label="⬇ Download Full Report (DOCX)", data=doc_bytes, file_name=f"report_{patient_id}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    update_patient_history(patient_id, patient_name, patient_age, patient_gender, class_name, confidence)
    st.success("✅ Analysis complete and patient record saved to history!")
    st.markdown('</div>', unsafe_allow_html=True)

def update_patient_history(p_id, p_name, p_age, p_gender, diagnosis, confidence):
    """Appends a new record to the patient history in session state."""
    new_record = pd.DataFrame([{"Timestamp": datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S'),"Patient ID": p_id, "Name": p_name, "Age": p_age, "Gender": p_gender,"Diagnosis": diagnosis, "Confidence": f"{confidence:.2%}"}])
    if 'patient_data' not in st.session_state or st.session_state.patient_data.empty:
        st.session_state.patient_data = new_record
    else:
        st.session_state.patient_data = pd.concat([st.session_state.patient_data, new_record], ignore_index=True)

def render_history_page():
    # (This function is unchanged)
    st.markdown('<div class="custom-section-container">', unsafe_allow_html=True)
    st.markdown('<h2 class="custom-section-header">📜 Patient History Records</h2>', unsafe_allow_html=True)
    if 'patient_data' in st.session_state and not st.session_state.patient_data.empty:
        st.dataframe(st.session_state.patient_data, use_container_width=True)
        csv = st.session_state.patient_data.to_csv(index=False).encode('utf-8')
        st.download_button("⬇ Download History (CSV)", csv, "patient_history.csv", "text/csv")
    else:
        st.info("🤷‍♀ No patient records have been created yet.")
    st.markdown('</div>', unsafe_allow_html=True)

def render_about_page():
    # (This function is unchanged)
    st.markdown('<div class="custom-section-container">', unsafe_allow_html=True)
    st.markdown('<h2 class="custom-section-header">💡 About This Application</h2>', unsafe_allow_html=True)
    st.info("This AI-powered web application assists in the preliminary analysis of brain MRI images. It integrates patient data management, AI-based classification, mock tumor localization, and Grad-CAM visualization for model explainability.")
    st.subheader("Disclaimer")
    st.warning("🛑 *Important:* This application is for informational and educational purposes only and should *NOT* be used for actual medical diagnosis. AI predictions are not a substitute for professional medical advice.")
    st.markdown("</div>", unsafe_allow_html=True)

# --- NEW: Function to render the enhanced sidebar ---
def render_sidebar():
    """Renders the feature-rich sidebar."""
    st.sidebar.title("App Controls & Info")
    st.sidebar.markdown("---")

    # Section 1: Model Information
    st.sidebar.subheader("🤖 Model Information")
    st.sidebar.info(
        """
        This app uses a refined EfficientNetV2-B2 model, fine-tuned on the Brain Tumor MRI Dataset.
        """
    )
    st.sidebar.metric("Test Accuracy", "98.5%+") # Note: Update with your model's actual accuracy
    st.sidebar.write("*Input Image Size:* 224x224 pixels")
    st.sidebar.markdown("---")

    # Section 2: Session Summary
    st.sidebar.subheader("📊 Session Summary")
    if 'patient_data' in st.session_state and not st.session_state.patient_data.empty:
        df = st.session_state.patient_data
        total_analyses = len(df)
        tumor_count = len(df[df['Diagnosis'] != 'No Tumor'])
        no_tumor_count = total_analyses - tumor_count
        st.sidebar.metric("Total Analyses This Session", total_analyses)
        col1, col2 = st.sidebar.columns(2)
        col1.metric("Tumor Found", tumor_count)
        col2.metric("No Tumor", no_tumor_count)
    else:
        st.sidebar.write("No analyses performed yet in this session.")

    # Section 3: Data Management
    if st.sidebar.button("🗑 Clear Session History"):
        st.session_state.patient_data = pd.DataFrame(columns=["Timestamp", "Patient ID", "Name", "Age", "Gender", "Diagnosis", "Confidence"])
        st.toast("Patient history for this session has been cleared.", icon="✅")
        st.rerun()
    st.sidebar.markdown("---")

    # Section 4: Resource Links
    st.sidebar.subheader("🔗 Resource Links")
    st.sidebar.markdown(
        """
        - [Learn about Brain Tumors (NCI)](https://www.cancer.gov/types/brain)
        - [View the Dataset on Kaggle](https://www.kaggle.com/datasets/masoudnickparvar/brain-tumor-mri-dataset)
        - [Project Source Code (GitHub)](https://github.com/)
        """
    )

# ==============================================================================
# 7. MAIN APP EXECUTION
# ==============================================================================
def main():
    """Main function to run the Streamlit app."""
    st.set_page_config(page_title=Config.APP_TITLE, page_icon=Config.PAGE_ICON, layout="wide", initial_sidebar_state="expanded")
    load_css()
    st.title(f"{Config.PAGE_ICON} {Config.APP_TITLE}")

    if 'patient_data' not in st.session_state:
        st.session_state.patient_data = pd.DataFrame(columns=["Timestamp", "Patient ID", "Name", "Age", "Gender", "Diagnosis", "Confidence"])

    # Render the new, feature-rich sidebar
    render_sidebar()

    # Main page tabs
    tab1, tab2, tab3 = st.tabs(["📊 Patient Analysis", "📜 Patient History", "💡 About & Help"])
    with tab1:
        render_analysis_page()
    with tab2:
        render_history_page()
    with tab3:
        render_about_page()

if __name__ == "__main__":
    main()
